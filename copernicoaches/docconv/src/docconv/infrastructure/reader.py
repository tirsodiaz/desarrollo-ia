"""
Infrastructure layer for reading documents from files.

This module provides functions to read DOCX and Markdown files
and convert them to internal Document structures.
"""

import os
from typing import Optional
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.text.run import Run
from docx.table import Table as DocxTable

from ..domain.models import (
    Document,
    DocumentElement,
    Paragraph,
    Heading,
    DocumentList,
    ListItem,
    Table,
    TableRow,
    TableCell,
    Blockquote,
    CodeBlock,
    HorizontalRule,
    FormattedText,
    TextFormat,
    Link,
    ConversionWarning,
)
from ..domain.md_to_docx import markdown_to_docx


def read_docx(file_path: str) -> tuple[Optional[Document], list[ConversionWarning]]:
    """
    Read a DOCX file and convert it to a Document structure.

    Args:
        file_path: Path to the .docx file

    Returns:
        Tuple of (document, warnings) or (None, warnings) if failed
    """
    warnings: list[ConversionWarning] = []

    if not os.path.exists(file_path):
        warnings.append(ConversionWarning(
            element_type="file",
            description=f"File does not exist: {file_path}",
            details={"file_path": file_path}
        ))
        return None, warnings

    try:
        docx_doc = DocxDocument(file_path)
    except Exception as e:
        warnings.append(ConversionWarning(
            element_type="file",
            description=f"Failed to read DOCX file: {str(e)}",
            details={"file_path": file_path, "error": str(e)}
        ))
        return None, warnings

    elements: list[DocumentElement] = []

    for para in docx_doc.paragraphs:
        if not para.text.strip():
            continue

        # Check if it's a heading
        style_name = para.style.name.lower()
        if 'heading' in style_name:
            level = 1
            if 'heading 1' in style_name:
                level = 1
            elif 'heading 2' in style_name:
                level = 2
            elif 'heading 3' in style_name:
                level = 3
            elif 'heading 4' in style_name:
                level = 4
            elif 'heading 5' in style_name:
                level = 5
            elif 'heading 6' in style_name:
                level = 6

            content = _parse_runs_to_formatted_text(para.runs)
            elements.append(Heading(level=level, content=content))
        else:
            # Regular paragraph
            content = _parse_runs_to_formatted_text(para.runs)
            elements.append(Paragraph(content=content))

    # Handle tables
    for table in docx_doc.tables:
        table_rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_content = []
                for para in cell.paragraphs:
                    cell_content.extend(_parse_runs_to_formatted_text(para.runs))
                cells.append(TableCell(content=cell_content))
            table_rows.append(TableRow(cells=cells))

        elements.append(Table(rows=table_rows))

    # Note: Lists, blockquotes, code blocks are not easily distinguishable in python-docx
    # without more complex parsing. For now, we'll treat them as paragraphs and add warnings.

    if not elements:
        warnings.append(ConversionWarning(
            element_type="document",
            description="No content found in DOCX file",
            details={"file_path": file_path}
        ))

    document = Document(elements=elements)
    return document, warnings


def _parse_runs_to_formatted_text(runs: list[Run]) -> list[FormattedText | str]:
    """Parse DOCX runs into formatted text elements."""
    elements = []
    for run in runs:
        text = run.text
        if not text:
            continue

        formats = []
        if run.bold:
            formats.append(TextFormat.BOLD)
        if run.italic:
            formats.append(TextFormat.ITALIC)

        if formats:
            elements.append(FormattedText(text=text, formats=formats))
        else:
            elements.append(text)

    return elements


def read_markdown(file_path: str) -> tuple[Optional[Document], list[ConversionWarning]]:
    """
    Read a Markdown file and convert it to a Document structure.

    Args:
        file_path: Path to the .md file

    Returns:
        Tuple of (document, warnings) or (None, warnings) if failed
    """
    warnings: list[ConversionWarning] = []

    if not os.path.exists(file_path):
        warnings.append(ConversionWarning(
            element_type="file",
            description=f"File does not exist: {file_path}",
            details={"file_path": file_path}
        ))
        return None, warnings

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
    except Exception as e:
        warnings.append(ConversionWarning(
            element_type="file",
            description=f"Failed to read Markdown file: {str(e)}",
            details={"file_path": file_path, "error": str(e)}
        ))
        return None, warnings

    document, parse_warnings = markdown_to_docx(markdown_text)
    warnings.extend(parse_warnings)

    return document, warnings