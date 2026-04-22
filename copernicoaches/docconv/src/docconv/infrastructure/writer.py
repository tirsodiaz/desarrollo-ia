"""
Infrastructure layer for writing documents to files.

This module provides functions to write Document structures
to DOCX and Markdown files.
"""

import os
from typing import Optional
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from ..domain.models import (
    Document,
    DocumentElement,
    Paragraph,
    Heading,
    DocumentList,
    ListItem,
    Table,
    TableRow,
    TableCell,
    Blockquote,
    CodeBlock,
    HorizontalRule,
    FormattedText,
    TextFormat,
    Link,
    ConversionWarning,
)
from ..domain.docx_to_md import docx_to_markdown


def write_docx(document: Document, file_path: str) -> list[ConversionWarning]:
    """
    Write a Document structure to a DOCX file.

    Args:
        document: The document to write
        file_path: Path where to save the .docx file

    Returns:
        List of warnings generated during writing
    """
    warnings: list[ConversionWarning] = []

    try:
        docx_doc = DocxDocument()

        for element in document.elements:
            if isinstance(element, Heading):
                _add_heading_to_docx(docx_doc, element)
            elif isinstance(element, Paragraph):
                _add_paragraph_to_docx(docx_doc, element)
            elif isinstance(element, DocumentList):
                _add_list_to_docx(docx_doc, element)
            elif isinstance(element, Table):
                _add_table_to_docx(docx_doc, element)
            elif isinstance(element, Blockquote):
                _add_blockquote_to_docx(docx_doc, element)
            elif isinstance(element, CodeBlock):
                _add_code_block_to_docx(docx_doc, element)
            elif isinstance(element, HorizontalRule):
                _add_horizontal_rule_to_docx(docx_doc)
            else:
                warnings.append(ConversionWarning(
                    element_type=type(element).__name__,
                    description=f"Unsupported element type for DOCX writing: {type(element).__name__}",
                    details={"element": str(element)}
                ))

        # Ensure directory exists
        dir_path = os.path.dirname(file_path)
        if dir_path:  # Only create directory if there's a directory path
            os.makedirs(dir_path, exist_ok=True)

        docx_doc.save(file_path)

    except Exception as e:
        warnings.append(ConversionWarning(
            element_type="file",
            description=f"Failed to write DOCX file: {str(e)}",
            details={"file_path": file_path, "error": str(e)}
        ))

    return warnings


def _add_heading_to_docx(docx_doc: DocxDocument, heading: Heading):
    """Add a heading to the DOCX document."""
    content = _format_text_elements(heading.content)
    para = docx_doc.add_heading(content, level=heading.level)


def _add_paragraph_to_docx(docx_doc: DocxDocument, paragraph: Paragraph):
    """Add a paragraph to the DOCX document."""
    content = _format_text_elements(paragraph.content)
    para = docx_doc.add_paragraph(content)


def _add_list_to_docx(docx_doc: DocxDocument, list_obj: DocumentList):
    """Add a list to the DOCX document."""
    for i, item in enumerate(list_obj.items, 1):
        content = _format_text_elements(item.content)
        if list_obj.ordered:
            prefix = f"{i}. "
        else:
            prefix = "- "
        docx_doc.add_paragraph(prefix + content)


def _add_table_to_docx(docx_doc: DocxDocument, table: Table):
    """Add a table to the DOCX document."""
    if not table.rows:
        return

    num_cols = len(table.rows[0].cells)
    docx_table = docx_doc.add_table(rows=len(table.rows), cols=num_cols)

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell_content = _format_text_elements(cell.content)
            docx_table.cell(i, j).text = cell_content


def _add_blockquote_to_docx(docx_doc: DocxDocument, blockquote: Blockquote):
    """Add a blockquote to the DOCX document."""
    content = _format_text_elements(blockquote.content)
    para = docx_doc.add_paragraph(content)
    # Try to set quote style, but don't fail if it doesn't exist
    try:
        para.style = 'Quote'
    except:
        pass  # Use default style


def _add_code_block_to_docx(docx_doc: DocxDocument, code_block: CodeBlock):
    """Add a code block to the DOCX document."""
    para = docx_doc.add_paragraph(code_block.code)
    # Try to set code style, but don't fail if it doesn't exist
    try:
        para.style = 'Code'
    except:
        pass  # Use default style


def _add_horizontal_rule_to_docx(docx_doc: DocxDocument):
    """Add a horizontal rule to the DOCX document."""
    para = docx_doc.add_paragraph()
    para.add_run().add_break()  # Simple line break as placeholder


def _format_text_elements(elements: list[FormattedText | str]) -> str:
    """Format text elements into a single string for DOCX."""
    # For simplicity, we'll concatenate and ignore formatting for now
    # Full formatting support would require more complex run manipulation
    parts = []
    for elem in elements:
        if isinstance(elem, str):
            parts.append(elem)
        else:
            parts.append(elem.text)
    return ''.join(parts)


def write_markdown(document: Document, file_path: str) -> list[ConversionWarning]:
    """
    Write a Document structure to a Markdown file.

    Args:
        document: The document to write
        file_path: Path where to save the .md file

    Returns:
        List of warnings generated during writing
    """
    warnings: list[ConversionWarning] = []

    try:
        markdown_text, conversion_warnings = docx_to_markdown(document)
        warnings.extend(conversion_warnings)

        # Ensure directory exists
        dir_path = os.path.dirname(file_path)
        if dir_path:  # Only create directory if there's a directory path
            os.makedirs(dir_path, exist_ok=True)

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)

    except Exception as e:
        warnings.append(ConversionWarning(
            element_type="file",
            description=f"Failed to write Markdown file: {str(e)}",
            details={"file_path": file_path, "error": str(e)}
        ))

    return warnings