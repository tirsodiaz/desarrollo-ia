import io
from pathlib import Path

import pytest
from docx import Document

from jano.infrastructure.reader import read_docx, read_md
from jano.infrastructure.writer import write_docx, write_md


def test_write_and_read_md(tmp_path):
    path = str(tmp_path / "test.md")
    write_md("# Hello\n", path)
    content = read_md(path)
    assert content == "# Hello\n"


def test_write_and_read_docx(tmp_path):
    path = str(tmp_path / "test.docx")
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("Test", level=1)
    doc.save(buf)
    write_docx(buf.getvalue(), path)
    loaded = read_docx(path)
    assert any(p.text == "Test" for p in loaded.paragraphs)


def test_read_md_missing_file():
    with pytest.raises(FileNotFoundError):
        read_md("/nonexistent/path/file.md")


def test_read_docx_missing_file():
    with pytest.raises(FileNotFoundError):
        read_docx("/nonexistent/path/file.docx")


def test_read_md_wrong_extension(tmp_path):
    bad = tmp_path / "file.txt"
    bad.write_text("hello")
    with pytest.raises(ValueError):
        read_md(str(bad))


def test_read_docx_wrong_extension(tmp_path):
    bad = tmp_path / "file.txt"
    bad.write_bytes(b"fake")
    with pytest.raises(ValueError):
        read_docx(str(bad))


def test_write_md_creates_parent_dirs(tmp_path):
    path = str(tmp_path / "sub" / "dir" / "out.md")
    write_md("content", path)
    assert Path(path).exists()


def test_write_docx_creates_parent_dirs(tmp_path):
    path = str(tmp_path / "sub" / "dir" / "out.docx")
    write_docx(b"\x00" * 4, path)
    assert Path(path).exists()
