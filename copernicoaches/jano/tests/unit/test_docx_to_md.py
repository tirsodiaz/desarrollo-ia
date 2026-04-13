from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from jano.domain.docx_to_md import convert_docx_to_md


def _doc_with_heading(level: int, text: str) -> Document:
    doc = Document()
    doc.add_heading(text, level=level)
    return doc


def _doc_with_paragraph(
    text: str, bold: bool = False, italic: bool = False
) -> Document:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return doc


def test_heading_level_1():
    result = convert_docx_to_md(_doc_with_heading(1, "Title"))
    assert result.content.startswith("# Title")


def test_heading_level_2():
    result = convert_docx_to_md(_doc_with_heading(2, "Section"))
    assert "## Section" in result.content


def test_heading_level_3():
    result = convert_docx_to_md(_doc_with_heading(3, "Subsection"))
    assert "### Subsection" in result.content


def test_bold_run():
    result = convert_docx_to_md(_doc_with_paragraph("hello", bold=True))
    assert "**hello**" in result.content


def test_italic_run():
    result = convert_docx_to_md(_doc_with_paragraph("world", italic=True))
    assert "*world*" in result.content


def test_plain_paragraph():
    result = convert_docx_to_md(_doc_with_paragraph("plain text"))
    assert "plain text" in result.content


def test_bullet_list():
    doc = Document()
    doc.add_paragraph("item one", style="List Bullet")
    result = convert_docx_to_md(doc)
    assert "- item one" in result.content


def test_table_gfm():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    result = convert_docx_to_md(doc)
    assert "| A | B |" in result.content
    assert "| --- | --- |" in result.content
    assert "| 1 | 2 |" in result.content


def test_image_produces_warning():
    doc = Document()
    # Add a paragraph and inject a fake blip element to simulate image
    p = doc.add_paragraph()
    r = p.add_run()
    drawing = OxmlElement("w:drawing")
    inline_el = OxmlElement("wp:inline")
    graphic = OxmlElement("a:graphic")
    graphicData = OxmlElement("a:graphicData")
    pic = OxmlElement("pic:pic")
    blipFill = OxmlElement("pic:blipFill")
    blip = OxmlElement("a:blip")
    blip.set(qn("r:embed"), "rId999")
    blipFill.append(blip)
    pic.append(blipFill)
    graphicData.append(pic)
    graphic.append(graphicData)
    inline_el.append(graphic)
    drawing.append(inline_el)
    r._r.append(drawing)

    result = convert_docx_to_md(doc)
    assert any(w.element_type == "image" for w in result.warnings)
    assert "[image]" in result.content


def test_no_warnings_for_plain_content():
    doc = Document()
    doc.add_heading("Clean", level=1)
    doc.add_paragraph("No issues here.")
    result = convert_docx_to_md(doc)
    assert result.warnings == []
