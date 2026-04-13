import io

from docx import Document

from jano.domain.md_to_docx import convert_md_to_docx


def _load(result) -> Document:
    assert isinstance(result.content, bytes)
    return Document(io.BytesIO(result.content))


def test_heading_1():
    result = convert_md_to_docx("# My Title\n")
    doc = _load(result)
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert any("My Title" in h.text for h in headings)


def test_heading_2():
    result = convert_md_to_docx("## Section\n")
    doc = _load(result)
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert any("Section" in h.text for h in headings)


def test_bold_inline():
    result = convert_md_to_docx("Hello **world** end\n")
    doc = _load(result)
    texts = [r.text for p in doc.paragraphs for r in p.runs]
    assert "world" in texts
    bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
    assert any("world" in r.text for r in bold_runs)


def test_italic_inline():
    result = convert_md_to_docx("Hello *world* end\n")
    doc = _load(result)
    italic_runs = [r for p in doc.paragraphs for r in p.runs if r.italic]
    assert any("world" in r.text for r in italic_runs)


def test_plain_paragraph():
    result = convert_md_to_docx("Just a paragraph.\n")
    doc = _load(result)
    texts = [p.text for p in doc.paragraphs]
    assert any("Just a paragraph." in t for t in texts)


def test_table_dimensions():
    md = "| A | B |\n|---|---|\n| 1 | 2 |\n"
    result = convert_md_to_docx(md)
    doc = _load(result)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 2
    assert len(doc.tables[0].columns) == 2


def test_code_block_monospace():
    md = "```\ncode here\n```\n"
    result = convert_md_to_docx(md)
    doc = _load(result)
    code_runs = [
        r
        for p in doc.paragraphs
        for r in p.runs
        if r.font.name and "Courier" in r.font.name
    ]
    assert code_runs


def test_raw_html_warning():
    md = "<div>some html</div>\n"
    result = convert_md_to_docx(md)
    assert any(w.element_type == "raw_html" for w in result.warnings)


def test_bullet_list():
    md = "- item one\n- item two\n"
    result = convert_md_to_docx(md)
    doc = _load(result)
    list_paras = [p for p in doc.paragraphs if "List" in p.style.name]
    assert len(list_paras) >= 2


def test_numbered_list():
    md = "1. first\n2. second\n"
    result = convert_md_to_docx(md)
    doc = _load(result)
    list_paras = [p for p in doc.paragraphs if "List" in p.style.name]
    assert len(list_paras) >= 2


def test_no_warnings_for_clean_markdown():
    md = "# Title\n\nParagraph.\n\n- item\n"
    result = convert_md_to_docx(md)
    assert result.warnings == []
