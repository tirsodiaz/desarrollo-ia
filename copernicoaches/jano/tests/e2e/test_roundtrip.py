"""Roundtrip stability test: convert multiple times, check no degradation."""

from pathlib import Path

from docx import Document

from jano.application.convert import convert_file

TEST_CASE = Path(__file__).parent.parent.parent / "test-case"
DOCX_FILE = str(TEST_CASE / "Que es un LLM.docx")
MD_FILE = str(TEST_CASE / "Que es un LLM.md")

# Content that MUST be present in every MD output
REQUIRED_MD_FRAGMENTS = [
    "LLM",  # heading content
    "##",  # at least one h2
    "- ",  # at least one bullet
    "|",  # table present
]


def _check_md(content: str, cycle: int) -> None:
    for fragment in REQUIRED_MD_FRAGMENTS:
        assert fragment in content, (
            f"Cycle {cycle}: expected {fragment!r} in MD output but not found.\n"
            f"Content (first 500 chars):\n{content[:500]}"
        )


def test_roundtrip_docx_start(tmp_path):
    """Start from DOCX, cycle 3 times through MD and back."""
    current_docx = DOCX_FILE
    for cycle in range(1, 4):
        md_out = str(tmp_path / f"cycle{cycle}.md")
        docx_out = str(tmp_path / f"cycle{cycle}.docx")

        convert_file(current_docx, md_out)
        md_content = Path(md_out).read_text(encoding="utf-8")
        _check_md(md_content, cycle)

        convert_file(md_out, docx_out)
        doc = Document(docx_out)
        headings = [p for p in doc.paragraphs if "Heading" in p.style.name]
        assert headings, f"Cycle {cycle}: no headings found in DOCX"

        current_docx = docx_out


def test_roundtrip_md_start(tmp_path):
    """Start from MD, cycle 3 times through DOCX and back."""
    current_md = MD_FILE
    for cycle in range(1, 4):
        docx_out = str(tmp_path / f"cycle{cycle}.docx")
        md_out = str(tmp_path / f"cycle{cycle}.md")

        convert_file(current_md, docx_out)
        doc = Document(docx_out)
        headings = [p for p in doc.paragraphs if "Heading" in p.style.name]
        assert headings, f"Cycle {cycle}: no headings found in DOCX"

        convert_file(docx_out, md_out)
        md_content = Path(md_out).read_text(encoding="utf-8")
        _check_md(md_content, cycle)

        current_md = md_out
