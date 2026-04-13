from __future__ import annotations

import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from markdown_it import MarkdownIt

from jano.domain.models import ConversionResult, ConversionWarning


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _apply_inline(paragraph, inline_tokens: list[dict]) -> None:
    """Apply inline tokens (bold, italic, text, hyperlinks) to a paragraph."""
    bold = False
    italic = False
    link_url: str | None = None
    link_text_parts: list[str] = []
    collecting_link = False

    for tok in inline_tokens:
        t = tok["type"]
        if t == "strong_open":
            bold = True
        elif t == "strong_close":
            bold = False
        elif t == "em_open":
            italic = True
        elif t == "em_close":
            italic = False
        elif t == "link_open":
            link_url = tok.get("attrs", {}).get("href", "")
            collecting_link = True
            link_text_parts = []
        elif t == "link_close":
            combined = "".join(link_text_parts)
            if link_url:
                _add_hyperlink(paragraph, combined, link_url)
            else:
                run = paragraph.add_run(combined)
                run.bold = bold
                run.italic = italic
            collecting_link = False
            link_url = None
            link_text_parts = []
        elif t == "text" or t == "code_inline":
            text = tok.get("content", "")
            if collecting_link:
                link_text_parts.append(text)
            else:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                if t == "code_inline":
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
        elif t == "softbreak":
            if not collecting_link:
                paragraph.add_run(" ")
        elif t == "hardbreak":
            if not collecting_link:
                paragraph.add_run("\n")
        elif t == "html_inline":
            # Strip HTML tags, keep text
            text = re.sub(r"<[^>]+>", "", tok.get("content", ""))
            if text.strip():
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic


def _set_list_style(paragraph, style_name: str, level: int) -> None:
    """Apply list paragraph style with correct numbering level."""
    paragraph.style = style_name
    # Set list indentation level via pPr > numPr > ilvl
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numPr.append(ilvl)
    # numId=1 for bullets, numId=2 for numbers (Word default template)
    numId_elem = OxmlElement("w:numId")
    num_id = "2" if "Number" in style_name else "1"
    numId_elem.set(qn("w:val"), num_id)
    numPr.append(numId_elem)
    pPr.append(numPr)


def _tokens_to_flat(tokens) -> list[dict]:
    """Flatten markdown-it token list to simple dicts."""
    result = []
    for tok in tokens:
        d = {"type": tok.type, "tag": tok.tag, "content": tok.content}
        if tok.attrs:
            d["attrs"] = dict(tok.attrs)
        if tok.children:
            d["children"] = _tokens_to_flat(tok.children)
        d["markup"] = tok.markup
        d["level"] = tok.level
        result.append(d)
    return result


def convert_md_to_docx(markdown_text: str) -> ConversionResult:
    warnings: list[ConversionWarning] = []
    doc = Document()

    md = MarkdownIt().enable("table")
    tokens = _tokens_to_flat(md.parse(markdown_text))

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        t = tok["type"]

        if t == "heading_open":
            level = int(tok["tag"][1])  # h1 → 1
            inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            text = inline_tok["content"] if inline_tok else ""
            p = doc.add_heading(level=level)
            p.clear()
            if inline_tok and inline_tok.get("children"):
                _apply_inline(p, inline_tok["children"])
            else:
                p.add_run(text)
            i += 3  # heading_open, inline, heading_close
            continue

        if t == "paragraph_open":
            inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
            p = doc.add_paragraph(style="Normal")
            if inline_tok and inline_tok.get("children"):
                _apply_inline(p, inline_tok["children"])
            elif inline_tok:
                p.add_run(inline_tok["content"])
            i += 3
            continue

        if t == "bullet_list_open":
            i += 1
            list_level = 0
            while i < len(tokens) and tokens[i]["type"] != "bullet_list_close":
                lt = tokens[i]["type"]
                if lt == "list_item_open":
                    list_level = tokens[i]["level"] // 2
                    i += 1
                    continue
                if lt == "paragraph_open":
                    inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                    p = doc.add_paragraph(style="List Bullet")
                    _set_list_style(p, "List Bullet", list_level)
                    if inline_tok and inline_tok.get("children"):
                        _apply_inline(p, inline_tok["children"])
                    elif inline_tok:
                        p.add_run(inline_tok["content"])
                    i += 3
                    continue
                if lt in ("bullet_list_open",):
                    list_level += 1
                elif lt in ("bullet_list_close",) and list_level > 0:
                    list_level -= 1
                i += 1
            i += 1  # skip bullet_list_close
            continue

        if t == "ordered_list_open":
            i += 1
            list_level = 0
            while i < len(tokens) and tokens[i]["type"] != "ordered_list_close":
                lt = tokens[i]["type"]
                if lt == "list_item_open":
                    list_level = tokens[i]["level"] // 2
                    i += 1
                    continue
                if lt == "paragraph_open":
                    inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                    p = doc.add_paragraph(style="List Number")
                    _set_list_style(p, "List Number", list_level)
                    if inline_tok and inline_tok.get("children"):
                        _apply_inline(p, inline_tok["children"])
                    elif inline_tok:
                        p.add_run(inline_tok["content"])
                    i += 3
                    continue
                if lt in ("ordered_list_open",):
                    list_level += 1
                elif lt in ("ordered_list_close",) and list_level > 0:
                    list_level -= 1
                i += 1
            i += 1
            continue

        if t == "fence":
            code = tok["content"]
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(code.rstrip("\n"))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            i += 1
            continue

        if t == "code_block":
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(tok["content"].rstrip("\n"))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            i += 1
            continue

        if t == "blockquote_open":
            i += 1
            while i < len(tokens) and tokens[i]["type"] != "blockquote_close":
                lt = tokens[i]["type"]
                if lt == "paragraph_open":
                    inline_tok = tokens[i + 1] if i + 1 < len(tokens) else None
                    p = doc.add_paragraph(style="Normal")
                    # Indent as blockquote via paragraph formatting
                    from docx.shared import Inches

                    p.paragraph_format.left_indent = Inches(0.5)
                    run = p.add_run("| ")
                    run.bold = True
                    if inline_tok and inline_tok.get("children"):
                        _apply_inline(p, inline_tok["children"])
                    elif inline_tok:
                        p.add_run(inline_tok["content"])
                    i += 3
                    continue
                i += 1
            i += 1
            continue

        if t == "table_open":
            i += 1
            header_cells: list[str] = []
            body_rows: list[list[str]] = []
            in_head = False
            current_row: list[str] = []
            while i < len(tokens) and tokens[i]["type"] != "table_close":
                lt = tokens[i]["type"]
                if lt == "thead_open":
                    in_head = True
                elif lt == "thead_close":
                    in_head = False
                elif lt in ("tbody_open", "tbody_close"):
                    pass
                elif lt == "tr_open":
                    current_row = []
                elif lt == "tr_close":
                    if in_head:
                        header_cells = current_row
                    else:
                        body_rows.append(current_row)
                elif lt == "th_open" or lt == "td_open":
                    pass
                elif lt == "inline":
                    current_row.append(tokens[i]["content"])
                i += 1
            i += 1  # table_close

            # Build DOCX table
            all_rows = ([header_cells] if header_cells else []) + body_rows
            if all_rows:
                num_cols = max(len(r) for r in all_rows)
                table = doc.add_table(rows=len(all_rows), cols=num_cols)
                table.style = "Table Grid"
                for r_idx, row_data in enumerate(all_rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = cell_text
                    if r_idx == 0 and header_cells:
                        for cell in table.rows[0].cells:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            continue

        if t == "hr":
            p = doc.add_paragraph("—" * 20)
            i += 1
            continue

        if t == "html_block":
            text = re.sub(r"<[^>]+>", "", tok["content"]).strip()
            warnings.append(
                ConversionWarning(
                    element_type="raw_html",
                    description=(
                        "Raw HTML block converted to plain text:"
                        f" {tok['content'][:60]!r}"
                    ),
                )
            )
            if text:
                doc.add_paragraph(text)
            i += 1
            continue

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return ConversionResult(content=buf.getvalue(), warnings=warnings)
